import logging
from io import BytesIO
from PyPDF2 import PdfReader
from docx import Document

logger = logging.getLogger(__name__)


class DocumentParser:
    """Service for parsing different document types"""
    
    @staticmethod
    def parse_text_file(file_content: bytes) -> str:
        """
        Parse plain text file
        
        Args:
            file_content: File content as bytes
            
        Returns:
            Extracted text
        """
        try:
            logger.info("Parsing text file")
            logger.debug(f"Text file size: {len(file_content)} bytes")
            
            text = file_content.decode('utf-8').strip()
            
            if not text:
                logger.error("Could not extract any text from text file")
                raise ValueError("Could not extract any text from text file")
            
            logger.info(f"Text file parsed successfully, extracted text length: {len(text)} characters")
            return text
            
        except UnicodeDecodeError as e:
            logger.error(f"Unicode decode error in text file: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to decode text file: {str(e)}")
        except Exception as e:
            logger.error(f"Error parsing text file: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse text file: {str(e)}")
    
    @staticmethod
    def parse_pdf_file(file_content: bytes) -> str:
        """
        Parse PDF file and extract text
        
        Args:
            file_content: File content as bytes
            
        Returns:
            Extracted text
        """
        try:
            logger.info("Parsing PDF file")
            logger.debug(f"PDF file size: {len(file_content)} bytes")
            
            # Create BytesIO object from file content
            pdf_file = BytesIO(file_content)
            
            # Parse PDF
            pdf_reader = PdfReader(pdf_file)
            num_pages = len(pdf_reader.pages)
            logger.debug(f"PDF has {num_pages} pages")
            
            # Extract text from all pages
            text = []
            for page_num, page in enumerate(pdf_reader.pages):
                logger.debug(f"Extracting text from PDF page {page_num}...")
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
                    logger.debug(f"Page {page_num} extracted, text length: {len(page_text)} characters")
                else:
                    logger.warning(f"Could not extract text from PDF page {page_num}")
            
            if not text:
                logger.error("Could not extract any text from PDF")
                raise ValueError("Could not extract any text from PDF")
            
            full_text = '\n'.join(text).strip()
            logger.info(f"PDF file parsed successfully, total extracted text length: {len(full_text)} characters")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error parsing PDF file: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse PDF file: {str(e)}")
    
    @staticmethod
    def parse_docx_file(file_content: bytes) -> str:
        """
        Parse DOCX file and extract text
        
        Args:
            file_content: File content as bytes
            
        Returns:
            Extracted text
        """
        try:
            logger.info("Parsing DOCX file")
            logger.debug(f"DOCX file size: {len(file_content)} bytes")
            
            # Create BytesIO object from file content
            docx_file = BytesIO(file_content)
            
            # Parse DOCX
            doc = Document(docx_file)
            logger.debug(f"DOCX has {len(doc.paragraphs)} paragraphs")
            
            # Extract text from all paragraphs
            text = []
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text.append(paragraph.text)
                    logger.debug(f"Paragraph {para_num} extracted, text length: {len(paragraph.text)} characters")
            
            if not text:
                logger.error("Could not extract any text from DOCX")
                raise ValueError("Could not extract any text from DOCX")
            
            full_text = '\n'.join(text).strip()
            logger.info(f"DOCX file parsed successfully, total extracted text length: {len(full_text)} characters")
            
            return full_text
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {str(e)}", exc_info=True)
            raise ValueError(f"Failed to parse DOCX file: {str(e)}")
    
    @staticmethod
    def parse_file(file_content: bytes, file_type: str) -> str:
        """
        Parse file based on type
        
        Args:
            file_content: File content as bytes
            file_type: File type (txt, pdf, docx)
            
        Returns:
            Extracted text
        """
        logger.info(f"Parsing file, type: {file_type}, size: {len(file_content)} bytes")
        
        file_type = file_type.lower().strip('.')
        
        logger.debug(f"Normalized file type: {file_type}")
        
        if file_type == 'txt':
            logger.debug("Using text parser")
            return DocumentParser.parse_text_file(file_content)
        elif file_type == 'pdf':
            logger.debug("Using PDF parser")
            return DocumentParser.parse_pdf_file(file_content)
        elif file_type in ['docx', 'doc']:
            logger.debug("Using DOCX parser")
            return DocumentParser.parse_docx_file(file_content)
        else:
            logger.error(f"Unsupported file type: {file_type}")
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def get_file_extension(filename: str) -> str:
        """
        Get file extension from filename
        
        Args:
            filename: Filename
            
        Returns:
            File extension without dot
        """
        logger.debug(f"Extracting file extension from filename: {filename}")
        
        if '.' not in filename:
            logger.error(f"Filename has no extension: {filename}")
            raise ValueError("Filename has no extension")
        
        extension = filename.rsplit('.', 1)[-1]
        logger.debug(f"Extracted file extension: {extension}")
        
        return extension
